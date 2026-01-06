from io import BytesIO
import PyPDF2
from docx import Document
from typing import List, Dict
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles processing of different document types (PDF, DOCX, TXT)"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_file(self, uploaded_file) -> List[Dict[str, str]]:
        """Process an uploaded file and return text chunks with metadata"""
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            
            if file_extension == 'pdf':
                text = self._extract_from_pdf(uploaded_file)
            elif file_extension == 'docx':
                text = self._extract_from_docx(uploaded_file)
            elif file_extension == 'txt':
                text = self._extract_from_txt(uploaded_file)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            if not text.strip():
                logger.warning(f"No text extracted from {uploaded_file.name}")
                return []
            
            # Split text into chunks
            chunks = self._split_text_into_chunks(text)
            
            # Create chunks with metadata
            processed_chunks = []
            for i, chunk in enumerate(chunks):
                processed_chunks.append({
                    "content": chunk,
                    "metadata": {
                        "source": uploaded_file.name,
                        "chunk_id": i,
                        "file_type": file_extension
                    }
                })
            
            logger.info(f"Processed {uploaded_file.name}: {len(processed_chunks)} chunks")
            return processed_chunks
            
        except Exception as e:
            logger.error(f"Error processing {uploaded_file.name}: {str(e)}")
            return []
    
    def _extract_from_pdf(self, uploaded_file) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.read()))
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    logger.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            return text
            
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_from_docx(self, uploaded_file) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(BytesIO(uploaded_file.read()))
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return text
            
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def _extract_from_txt(self, uploaded_file) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    uploaded_file.seek(0)  # Reset file pointer
                    text = uploaded_file.read().decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with any supported encoding")
            
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
    
    def _split_text_into_chunks(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # If we're not at the end, try to break at a sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                last_period = text.rfind('.', start, end)
                last_exclamation = text.rfind('!', start, end)
                last_question = text.rfind('?', start, end)
                
                sentence_end = max(last_period, last_exclamation, last_question)
                
                if sentence_end > start:
                    end = sentence_end + 1
                else:
                    # Look for word boundary
                    last_space = text.rfind(' ', start, end)
                    if last_space > start:
                        end = last_space
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            if start <= 0:
                break
        
        return chunks
